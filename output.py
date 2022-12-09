import deepl
import docx
import glob
import os
import shutil
import pyperclip
import time
import keyboard
import subprocess

def copy_text():
    while True:
        flag = False
        title = input('セクション名を入力してください：')

        title = title.replace(' ', '_')
        title = title + '.txt'

        subprocess.Popen(['code', title], shell=True)

        while True:
            f = open(title, 'a', encoding='UTF-8')
            originalText = pyperclip.paste()
            while True:
                currentText = pyperclip.paste()
                if currentText != originalText:
                    currentText = currentText.replace('\r\n', ' ')
                    f.write(currentText)
                    f.close()
                    break
                time.sleep(0.01)
                if keyboard.is_pressed('enter'):
                # if True:
                    f.write('\n')
                    f.close
                    time.sleep(0.5)
                    break
                if keyboard.is_pressed('escape'):
                    f.close
                    flag = True
                    break
            if flag == True:
                break
            
        ans = input('セクションを追加しますか？(y/n):')
        if ans == 'n':
            break

def translation(API_KEY, source_lang, target_lang):
# イニシャライズ
    translator = deepl.Translator(API_KEY)

    if glob.glob("*.docx"):
        doc = docx.Document(glob.glob("*.docx")[0])
    else:
        doc = docx.Document()

    if not glob.glob("Archive"):
        os.mkdir("Archive")

    txts = glob.glob("*.txt")

    for txt in txts:
        doc.add_paragraph(txt.replace("_", " ").replace(".txt", ""))
        doc.paragraphs[-1].runs[0].bold = True
        doc.paragraphs[-1].runs[0].font.size = docx.shared.Pt(16)

        f = open(txt, 'r', encoding="utf-8")

        for text in f.readlines():
            # 翻訳を実行
            result = translator.translate_text(text, source_lang=source_lang, target_lang=target_lang)
            # 段落の末尾に追加
            doc.add_paragraph(result.text)

        f.close()
        shutil.move(txt, "./Archive/")

    doc.save("Sample.docx")


if __name__ == '__main__':
    API_KEY = 'af8b9802-207d-2d98-20e7-05985839ff75:fx' # 自身の API キーを指定
    source_lang = 'EN'
    target_lang = 'JA'
    copy_text()
    ans = input('翻訳を行いますか？(y/n):')
    if ans == 'y':
        translation(API_KEY, source_lang, target_lang)
    else:
        pass